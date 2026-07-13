"""Exportação da análise de credores para um relatório Word (.docx).

Estrutura: resumo executivo, KPIs, tabela de credores, gráfico por classe,
ranking dos maiores credores, análise de concentração, simulações de
aquisição de quórum e conclusões (cenários técnicos, sem aconselhamento
jurídico ou financeiro).
"""

from __future__ import annotations

import io
from datetime import date
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # backend sem interface gráfica, necessário para servidores/Streamlit
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches, Pt, RGBColor

from config import CLASSE_COR_PADRAO, CLASSE_CORES, CORES, LOGO_PATH, NOME_EMPRESA, NOME_SISTEMA, configurar_logging
from src import analise_quorum, estrategia
from src.models import ResultadoExtracao
from src.utils import formatar_moeda, formatar_percentual

logger = configurar_logging()

_COR_PRIMARIA_RGB = RGBColor(0x24, 0x22, 0x88)


def _cor_hex_para_rgb(cor_hex: str) -> tuple[float, float, float]:
    cor_hex = cor_hex.lstrip("#")
    return tuple(int(cor_hex[i : i + 2], 16) / 255 for i in (0, 2, 4))


def _adicionar_cabecalho(doc: Document, nome_arquivo_pdf: str) -> None:
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(1.5))

    titulo = doc.add_heading(NOME_SISTEMA, level=0)
    titulo.runs[0].font.color.rgb = _COR_PRIMARIA_RGB

    subtitulo = doc.add_paragraph()
    subtitulo.add_run(f"{NOME_EMPRESA} — Relatório de Análise de Credores").bold = True
    subtitulo.add_run(f"\nDocumento analisado: {nome_arquivo_pdf}")
    subtitulo.add_run(f"\nData de geração: {date.today().strftime('%d/%m/%Y')}")

    aviso = doc.add_paragraph()
    run = aviso.add_run(
        "As análises a seguir são cenários técnicos baseados exclusivamente nos dados "
        "extraídos do documento. Não constituem aconselhamento jurídico ou financeiro."
    )
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()


def _adicionar_resumo_executivo(doc: Document, resultado: ResultadoExtracao, resumo_executivo: str | None) -> None:
    doc.add_heading("Resumo Executivo", level=1)

    valor_total = analise_quorum.valor_total_passivo(resultado.credores)
    tabela_kpi = doc.add_table(rows=1, cols=4)
    tabela_kpi.style = "Light Grid Accent 1"
    cabecalho = tabela_kpi.rows[0].cells
    cabecalho[0].text = "Total de Credores"
    cabecalho[1].text = "Valor Total do Passivo"
    cabecalho[2].text = "Pendentes de Revisão"
    cabecalho[3].text = "Páginas via OCR"

    linha = tabela_kpi.add_row().cells
    linha[0].text = str(resultado.total_credores)
    linha[1].text = formatar_moeda(valor_total)
    linha[2].text = str(len(resultado.credores_para_revisar) + len(resultado.credores_com_erro))
    linha[3].text = str(len(resultado.paginas_ocr))

    doc.add_paragraph()

    if resumo_executivo:
        doc.add_paragraph(resumo_executivo)
    else:
        doc.add_paragraph(
            "Resumo executivo não gerado por IA para este relatório. "
            "Os indicadores acima e as seções seguintes refletem exclusivamente "
            "os dados extraídos do documento."
        )


def _adicionar_avisos_reconciliacao(doc: Document, resultado: ResultadoExtracao) -> None:
    """Alerta quando os subtotais/total geral impressos no PDF não batem com o
    extraído (ferramentas de leitura de tabela podem perder linhas em quebras
    de página). Os dados extraídos nunca são alterados automaticamente.
    """
    if not resultado.avisos_reconciliacao:
        return

    doc.add_heading("Divergências de Reconciliação", level=1)
    aviso = doc.add_paragraph()
    run = aviso.add_run(
        "Os totais abaixo, impressos no próprio documento, não batem com a soma dos "
        "credores extraídos. Confira manualmente antes de usar os números para decisão."
    )
    run.bold = True
    for texto in resultado.avisos_reconciliacao:
        doc.add_paragraph(texto, style="List Bullet")


def _adicionar_tabela_dataframe(doc: Document, df: pd.DataFrame, colunas_moeda: set[str] = frozenset(), colunas_percentual: set[str] = frozenset()) -> None:
    if df.empty:
        doc.add_paragraph("Sem dados disponíveis para esta seção.")
        return

    tabela = doc.add_table(rows=1, cols=len(df.columns))
    tabela.style = "Light Grid Accent 1"
    for i, coluna in enumerate(df.columns):
        tabela.rows[0].cells[i].text = str(coluna)

    for _, linha_df in df.iterrows():
        celulas = tabela.add_row().cells
        for i, coluna in enumerate(df.columns):
            valor = linha_df[coluna]
            if coluna in colunas_moeda:
                texto = formatar_moeda(valor)
            elif coluna in colunas_percentual:
                texto = formatar_percentual(valor)
            else:
                texto = str(valor)
            celulas[i].text = texto


def _grafico_valor_por_classe(df_resumo_classe: pd.DataFrame) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    cores = [
        _cor_hex_para_rgb(CLASSE_CORES.get(classe, CLASSE_COR_PADRAO))
        for classe in df_resumo_classe["Classe"]
    ]
    ax.bar(df_resumo_classe["Classe"], df_resumo_classe["Valor Total"], color=cores)
    ax.set_ylabel("Valor Total (R$)")
    ax.set_title("Valor Total por Classe")
    plt.xticks(rotation=20, ha="right", fontsize=8)
    fig.tight_layout()

    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _adicionar_conclusoes(doc: Document) -> None:
    doc.add_heading("Conclusões", level=1)
    doc.add_paragraph(
        "Este relatório apresenta cenários técnicos de concentração de crédito e simulações "
        "de formação de quórum, construídos exclusivamente a partir dos dados extraídos do "
        "documento de credores. Registros marcados como pendentes de revisão devem ser "
        "conferidos manualmente antes de qualquer decisão. Este material não constitui "
        "aconselhamento jurídico ou financeiro e não substitui a análise de um profissional "
        "habilitado."
    )


def exportar_word(
    resultado: ResultadoExtracao,
    caminho_saida: str | Path,
    resumo_executivo: str | None = None,
    top_n_ranking: int = 15,
) -> Path:
    """Gera um relatório .docx completo a partir do resultado da extração."""
    caminho_saida = Path(caminho_saida)
    caminho_saida.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    _adicionar_cabecalho(doc, resultado.arquivo_nome)
    _adicionar_resumo_executivo(doc, resultado, resumo_executivo)
    _adicionar_avisos_reconciliacao(doc, resultado)

    df_resumo_classe = analise_quorum.resumo_por_classe(resultado.credores)
    doc.add_heading("Resumo por Classe", level=1)
    _adicionar_tabela_dataframe(
        doc, df_resumo_classe, colunas_moeda={"Valor Total"}, colunas_percentual={"% do Passivo Total"}
    )

    if not df_resumo_classe.empty:
        doc.add_paragraph()
        doc.add_picture(_grafico_valor_por_classe(df_resumo_classe), width=Inches(6))

    doc.add_heading("Ranking dos Maiores Credores", level=1)
    df_ranking = analise_quorum.ranking_maiores_credores(resultado.credores, top_n=top_n_ranking)
    colunas_ranking = [c for c in ["Ranking", "Nome", "Documento", "Classe", "Valor", "% do Passivo Total", "Participação Acumulada"] if c in df_ranking.columns]
    _adicionar_tabela_dataframe(
        doc,
        df_ranking[colunas_ranking] if not df_ranking.empty else df_ranking,
        colunas_moeda={"Valor"},
        colunas_percentual={"% do Passivo Total", "Participação Acumulada"},
    )

    doc.add_heading("Análise de Concentração", level=1)
    concentracao = estrategia.concentracao_votos(resultado.credores)
    texto_concentracao = (
        f"Os 1 maior(es) credor(es) concentram {formatar_percentual(concentracao.get('top_1', 0))} do passivo total. "
        f"Os 5 maiores concentram {formatar_percentual(concentracao.get('top_5', 0))}, "
        f"os 10 maiores {formatar_percentual(concentracao.get('top_10', 0))} "
        f"e os 20 maiores {formatar_percentual(concentracao.get('top_20', 0))}."
    )
    doc.add_paragraph(texto_concentracao)

    doc.add_heading("Simulações de Formação de Quórum", level=1)
    simulacoes = estrategia.simular_formacao_quorum(resultado.credores)
    df_simulacoes = estrategia.tabela_simulacoes(simulacoes)
    _adicionar_tabela_dataframe(
        doc,
        df_simulacoes,
        colunas_moeda={"Valor Alvo", "Valor a Adquirir"},
        colunas_percentual={"Quórum Alvo", "Quórum Atingido"},
    )

    _adicionar_conclusoes(doc)

    doc.save(caminho_saida)
    logger.info("Word exportado em '%s' (%d credores).", caminho_saida, resultado.total_credores)
    return caminho_saida
