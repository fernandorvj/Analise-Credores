"""Exportação para Word (.docx) dos resultados da Calculadora — Simulador de
Financiamento e Calculadora de VPL.

Reaproveita a capa, o sumário e o rodapé paginado já compartilhados em
`src/exportar_word_base.py` (usados também pelo módulo Petição Inicial) —
sem alterar aquele arquivo. Gráficos são renderizados com `matplotlib`
(mesmo padrão já usado em `src/exportar_word.py` para o relatório de
Credores: figura → PNG em memória → `doc.add_picture`).
"""

from __future__ import annotations

import io
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # backend sem interface gráfica, necessário para servidores/Streamlit
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches, Pt

from config import NOME_EMPRESA
from src import exportar_word_base as base
from src.calculadora.exportar_excel import _df_cronograma_financiamento, _df_fluxo
from src.calculadora.models import ResultadoFinanciamento, ResultadoVPL
from src.utils import formatar_moeda, formatar_percentual

_COR_GRAFICO = "#3633CC"
_COR_GRAFICO_2 = "#B87A12"


def _grafico_saldo_devedor(resultado: ResultadoFinanciamento) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    numeros = [p.numero for p in resultado.parcelas]
    saldos = [float(p.saldo_final) for p in resultado.parcelas]
    ax.plot(numeros, saldos, color=_COR_GRAFICO, linewidth=2)
    ax.fill_between(numeros, saldos, color=_COR_GRAFICO, alpha=0.12)
    ax.set_title("Evolução do Saldo Devedor")
    ax.set_xlabel("Parcela")
    ax.set_ylabel("Saldo (R$)")
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _grafico_composicao_parcela(resultado: ResultadoFinanciamento) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    numeros = [p.numero for p in resultado.parcelas]
    juros = [float(p.juros) for p in resultado.parcelas]
    amortizacao = [float(p.amortizacao) for p in resultado.parcelas]
    ax.bar(numeros, amortizacao, color=_COR_GRAFICO, label="Amortização")
    ax.bar(numeros, juros, bottom=amortizacao, color=_COR_GRAFICO_2, label="Juros")
    ax.set_title("Composição da Parcela: Juros x Amortização")
    ax.set_xlabel("Parcela")
    ax.set_ylabel("Valor (R$)")
    ax.legend()
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _grafico_fluxo_vpl(resultado: ResultadoVPL) -> io.BytesIO:
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    datas = [item.data for item in resultado.fluxo_descontado]
    nominal = [float(item.valor) for item in resultado.fluxo_descontado]
    presente = [float(item.valor_presente) if item.valor_presente is not None else 0.0 for item in resultado.fluxo_descontado]
    ax.plot(datas, nominal, marker="o", color=_COR_GRAFICO, label="Valor Nominal")
    ax.plot(datas, presente, marker="o", color=_COR_GRAFICO_2, label="Valor Presente")
    ax.set_title("Fluxo de Recebimentos: Nominal x Valor Presente")
    ax.set_ylabel("Valor (R$)")
    plt.xticks(rotation=25, ha="right", fontsize=8)
    ax.legend()
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def exportar_word_financiamento(resultado: ResultadoFinanciamento, caminho_saida: str | Path) -> Path:
    """Gera o relatório Word do Simulador de Financiamento: capa, sumário,
    resumo, entradas utilizadas, premissas, cronograma, gráficos, resultados
    e observações.
    """
    caminho_saida = Path(caminho_saida)
    p = resultado.parametros
    doc = Document()

    base.adicionar_capa(
        doc,
        nome_arquivo_pdf=f"Simulação de Financiamento ({p.sistema.value})",
        subtitulo_modulo="Calculadora — Simulador de Financiamento",
        texto_aviso=(
            "Esta simulação é um cenário técnico calculado a partir dos parâmetros informados pelo "
            "usuário — não constitui proposta, oferta de crédito ou aconselhamento financeiro."
        ),
        pagina_isolada=True,
    )
    base.adicionar_sumario(doc)

    doc.add_heading("Resumo", level=1)
    doc.add_paragraph(
        f"Financiamento de {formatar_moeda(float(p.valor_financiado))} pelo sistema {p.sistema.value}, "
        f"em {p.prazo} parcelas {p.periodicidade_parcela.value.lower()}s, com taxa de "
        f"{formatar_percentual(float(p.taxa))} {p.periodicidade_taxa.value.lower()}. "
        f"Juros totais estimados em {formatar_moeda(float(resultado.juros_totais))}."
    )

    doc.add_heading("Entradas Utilizadas", level=1)
    df_entradas = pd.DataFrame(
        [
            {"Parâmetro": "Valor Financiado", "Valor": formatar_moeda(float(p.valor_financiado))},
            {"Parâmetro": "Valor de Entrada", "Valor": formatar_moeda(float(p.valor_entrada))},
            {"Parâmetro": "Taxa Informada", "Valor": f"{formatar_percentual(float(p.taxa))} {p.periodicidade_taxa.value}"},
            {"Parâmetro": "Periodicidade das Parcelas", "Valor": p.periodicidade_parcela.value},
            {"Parâmetro": "Prazo", "Valor": f"{p.prazo} parcelas"},
            {"Parâmetro": "Carência", "Valor": f"{p.carencia} parcelas"},
            {"Parâmetro": "Data Inicial", "Valor": p.data_inicial.strftime("%d/%m/%Y")},
            {"Parâmetro": "Sistema de Amortização", "Valor": p.sistema.value},
            {"Parâmetro": "Regime de Juros", "Valor": p.regime.value},
        ]
    )
    base.adicionar_tabela_dataframe(doc, df_entradas)

    doc.add_heading("Premissas", level=1)
    doc.add_paragraph(
        "O regime de juros (simples/compostos) se aplica à conversão da taxa informada entre "
        "periodicidades e à capitalização de juros durante a carência; o cálculo das parcelas "
        f"em si segue sempre a convenção padrão de mercado do sistema {p.sistema.value}. "
        "Todos os valores monetários são calculados com precisão decimal (sem erro de "
        "arredondamento de ponto flutuante) e a última parcela absorve eventual resíduo de "
        "centavos, fechando o saldo devedor exatamente em zero."
    )

    doc.add_heading("Cronograma", level=1)
    base.adicionar_tabela_dataframe(
        doc,
        _df_cronograma_financiamento(resultado),
        colunas_moeda={"Saldo Inicial", "Juros", "Amortização", "Valor Parcela", "Saldo Final"},
    )

    doc.add_heading("Gráficos", level=1)
    doc.add_picture(_grafico_saldo_devedor(resultado), width=Inches(6))
    doc.add_picture(_grafico_composicao_parcela(resultado), width=Inches(6))

    doc.add_heading("Resultados", level=1)
    df_resultados = pd.DataFrame(
        [
            {"Indicador": "Valor da Parcela Regular", "Valor": formatar_moeda(float(resultado.valor_parcela_regular or 0))},
            {"Indicador": "Juros Totais", "Valor": formatar_moeda(float(resultado.juros_totais))},
            {"Indicador": "Total Pago (incl. entrada)", "Valor": formatar_moeda(float(resultado.total_pago))},
            {"Indicador": "Taxa Periódica Efetiva", "Valor": formatar_percentual(float(resultado.taxa_periodica))},
        ]
    )
    base.adicionar_tabela_dataframe(doc, df_resultados)

    doc.add_heading("Observações", level=1)
    doc.add_paragraph(
        "Simulação gerada automaticamente pela Calculadora AMF3 Capital. Valores sujeitos a "
        "alteração conforme condições reais de contratação."
    )

    base.adicionar_rodape_paginacao(doc, f"{NOME_EMPRESA} — Simulador de Financiamento")
    doc.save(caminho_saida)
    return caminho_saida


def exportar_word_vpl(resultado: ResultadoVPL, caminho_saida: str | Path) -> Path:
    """Gera o relatório Word da Calculadora de VPL: capa, sumário, resumo,
    entradas utilizadas, premissas, fluxo, gráficos, indicadores e observações.
    """
    caminho_saida = Path(caminho_saida)
    p = resultado.parametros
    doc = Document()

    base.adicionar_capa(
        doc,
        nome_arquivo_pdf="Calculadora de VPL — Aquisição de Crédito",
        subtitulo_modulo="Calculadora — VPL de Aquisição de Créditos",
        texto_aviso=(
            "Esta análise é um cenário técnico calculado a partir dos parâmetros informados e da "
            "taxa de desconto vigente na data de geração — não constitui garantia de resultado, "
            "proposta de investimento ou aconselhamento financeiro."
        ),
        pagina_isolada=True,
    )
    base.adicionar_sumario(doc)

    doc.add_heading("Resumo", level=1)
    doc.add_paragraph(
        f"Aquisição de crédito de {formatar_moeda(float(p.valor_credito))} por "
        f"{formatar_moeda(float(p.valor_compra))} ({formatar_percentual(float(p.desagio))} de deságio), "
        f"descontado a {formatar_percentual(float(p.taxa_desconto_anual))} a.a. ({p.origem_taxa_desconto}). "
        f"VPL estimado em {formatar_moeda(float(resultado.vpl))}."
    )

    doc.add_heading("Entradas Utilizadas", level=1)
    df_entradas = pd.DataFrame(
        [
            {"Parâmetro": "Valor do Crédito", "Valor": formatar_moeda(float(p.valor_credito))},
            {"Parâmetro": "Valor de Compra", "Valor": formatar_moeda(float(p.valor_compra))},
            {"Parâmetro": "Deságio", "Valor": formatar_percentual(float(p.desagio))},
            {"Parâmetro": "Data Base", "Valor": p.data_base.strftime("%d/%m/%Y")},
            {"Parâmetro": "Taxa de Desconto (a.a.)", "Valor": formatar_percentual(float(p.taxa_desconto_anual))},
            {"Parâmetro": "Origem da Taxa de Desconto", "Valor": p.origem_taxa_desconto},
            {"Parâmetro": "Correção Monetária (a.a.)", "Valor": formatar_percentual(float(p.correcao_monetaria_anual))},
            {"Parâmetro": "Quantidade de Recebimentos", "Valor": str(len(p.fluxo_recebimentos))},
        ]
    )
    base.adicionar_tabela_dataframe(doc, df_entradas)

    doc.add_heading("Premissas", level=1)
    doc.add_paragraph(
        "VPL e TIR seguem a metodologia XNPV/XIRR (fluxos de caixa com datas irregulares, base de "
        "365 dias/ano) — a mesma convenção usada por planilhas financeiras (Excel/Google Sheets), "
        "mais adequada que um NPV de período fixo quando os recebimentos não são uniformemente "
        "espaçados. Valor Econômico é o valor presente apenas dos recebimentos esperados; VPL é o "
        "Valor Econômico menos o Valor de Compra."
    )

    doc.add_heading("Fluxo", level=1)
    base.adicionar_tabela_dataframe(
        doc,
        _df_fluxo(resultado.fluxo_descontado),
        colunas_moeda={"Valor", "Juros", "Amortização", "Saldo Devedor", "Valor Presente"},
    )

    doc.add_heading("Gráficos", level=1)
    doc.add_picture(_grafico_fluxo_vpl(resultado), width=Inches(6))

    doc.add_heading("Resultados e Indicadores", level=1)
    df_resultados = pd.DataFrame(
        [
            {"Indicador": "Valor Futuro", "Valor": formatar_moeda(float(resultado.valor_futuro))},
            {"Indicador": "Valor Econômico", "Valor": formatar_moeda(float(resultado.valor_economico))},
            {"Indicador": "VPL", "Valor": formatar_moeda(float(resultado.vpl))},
            {
                "Indicador": "TIR (a.a.)",
                "Valor": formatar_percentual(float(resultado.tir_anual)) if resultado.tir_anual is not None else "Não convergiu",
            },
            {
                "Indicador": "Taxa Efetiva (a.a.)",
                "Valor": (
                    formatar_percentual(float(resultado.taxa_efetiva_anual))
                    if resultado.taxa_efetiva_anual is not None
                    else "Não convergiu"
                ),
            },
            {"Indicador": "Payback", "Valor": resultado.payback_data.strftime("%d/%m/%Y") if resultado.payback_data else "Não atingido"},
            {"Indicador": "ROI", "Valor": formatar_percentual(float(resultado.roi)) if resultado.roi is not None else "-"},
            {
                "Indicador": "Rentabilidade",
                "Valor": formatar_percentual(float(resultado.rentabilidade)) if resultado.rentabilidade is not None else "-",
            },
            {"Indicador": "Margem", "Valor": formatar_percentual(float(resultado.margem)) if resultado.margem is not None else "-"},
            {"Indicador": "Spread (a.a.)", "Valor": formatar_percentual(float(resultado.spread)) if resultado.spread is not None else "-"},
        ]
    )
    base.adicionar_tabela_dataframe(doc, df_resultados)

    doc.add_heading("Observações", level=1)
    aviso_texto = (
        "TIR/Taxa Efetiva não convergiu para este fluxo — revise os valores informados."
        if resultado.tir_anual is None
        else "Análise gerada automaticamente pela Calculadora AMF3 Capital."
    )
    doc.add_paragraph(aviso_texto)

    base.adicionar_rodape_paginacao(doc, f"{NOME_EMPRESA} — Calculadora de VPL")
    doc.save(caminho_saida)
    return caminho_saida
