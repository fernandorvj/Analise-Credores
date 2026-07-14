"""Exportação do relatório do módulo Análise de Documentos para Word (.docx).

Reaproveita a marca visual e os helpers genéricos de
`src/exportar_word_base.py` (mesmo módulo usado pelos relatórios de
Credores, Petição Inicial e Precificação Inteligente), sem duplicar
capa/sumário/rodapé.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from config import NOME_EMPRESA, configurar_logging
from src import exportar_word_base as base
from src.models_analise_documentos import SECOES, AnaliseDocumento, ItemComContexto

logger = configurar_logging()

_TEXTO_AVISO_CAPA = (
    "Este relatório é gerado com apoio de Inteligência Artificial a partir do documento enviado, "
    "refletindo exclusivamente o que foi identificado no texto. Não constitui aconselhamento "
    "jurídico ou financeiro e não garante resultado."
)


def _renderizar_secao(doc: Document, titulo: str, valor: object) -> None:
    doc.add_heading(titulo, level=1)

    if isinstance(valor, list):
        if not valor:
            doc.add_paragraph("Não foi possível identificar itens para esta seção no documento.")
        elif isinstance(valor[0], ItemComContexto):
            df = pd.DataFrame([{"Item": i.item, "Contexto": i.contexto} for i in valor])
            base.adicionar_tabela_dataframe(doc, df)
        else:
            for item in valor:
                doc.add_paragraph(str(item), style="List Bullet")
        return

    doc.add_paragraph(str(valor) or "Não foi possível gerar esta seção automaticamente.")


def exportar_word_analise_documentos(analise: AnaliseDocumento, caminho_saida: str | Path) -> Path:
    """Gera o relatório .docx completo (capa + sumário + seções) a partir de
    uma `AnaliseDocumento` já preenchida (ver `src/ia.py`)."""
    caminho_saida = Path(caminho_saida)
    caminho_saida.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    base.adicionar_capa(
        doc,
        analise.arquivo_nome,
        subtitulo_modulo=f"Relatório de Análise de Documentos — {analise.tipo_origem}",
        texto_aviso=_TEXTO_AVISO_CAPA,
        pagina_isolada=True,
    )
    base.adicionar_sumario(doc)

    for titulo, chave in SECOES:
        _renderizar_secao(doc, titulo, getattr(analise, chave))

    if analise.avisos:
        doc.add_heading("Avisos sobre a Extração/Geração", level=1)
        for aviso in analise.avisos:
            doc.add_paragraph(aviso, style="List Bullet")

    base.adicionar_rodape_paginacao(doc, f"{NOME_EMPRESA} — Análise de Documentos")

    doc.save(caminho_saida)
    logger.info("Word da Análise de Documentos exportado em '%s' (%s).", caminho_saida, analise.arquivo_nome)
    return caminho_saida
