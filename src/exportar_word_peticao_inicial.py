"""Exportação do relatório do módulo Petição Inicial para Word (.docx).

Estrutura: capa, sumário, as 13 seções do relatório (nesta ordem — ver
`SECOES` em `src/models_peticao_inicial.py`) e um bloco final de avisos sobre
a extração/geração. Reaproveita a marca visual e os helpers genéricos de
`src/exportar_word_base.py` (mesmo módulo usado pelo relatório de Credores),
nunca duplicando a lógica de capa/tabela/rodapé.
"""

from __future__ import annotations

from pathlib import Path

import pandas as pd
from docx import Document

from config import NOME_EMPRESA, configurar_logging
from src import exportar_word_base as base
from src.models_peticao_inicial import (
    SECOES,
    DadosEmpresa,
    EventoCronologia,
    ItemComJustificativa,
    PassivoFiscal,
    RelatorioPeticaoInicial,
)

_ICONE_GRAU_ATENCAO = {"Baixo": "🟢", "Médio": "🟡", "Alto": "🔴"}

logger = configurar_logging()

_TEXTO_AVISO_CAPA = (
    "Este relatório é gerado com apoio de Inteligência Artificial a partir do documento "
    "enviado, refletindo exclusivamente o que foi identificado no texto. Não constitui "
    "aconselhamento jurídico ou financeiro e não garante resultado."
)


def _renderizar_dados_empresa(doc: Document, dados: DadosEmpresa) -> None:
    campos = [
        ("Razão Social", dados.razao_social),
        ("Nome Fantasia", dados.nome_fantasia),
        ("CNPJ", dados.cnpj),
        ("Segmento", dados.segmento),
        ("Atividade", dados.atividade),
        ("Grupo Econômico", dados.grupo_economico),
        ("Nº de Funcionários", dados.numero_funcionarios),
        ("Filiais", dados.filiais),
        ("Mercado de Atuação", dados.mercado_atuacao),
        *dados.outros_dados,
    ]
    base.adicionar_tabela_dataframe(doc, pd.DataFrame(campos, columns=["Campo", "Valor"]))


def _renderizar_passivo_fiscal(doc: Document, pf: PassivoFiscal) -> None:
    doc.add_heading("Situação Encontrada", level=2)
    df_situacao = pd.DataFrame(
        [
            ("Existe Passivo Fiscal?", pf.existe_passivo_fiscal),
            ("Existe Execução Fiscal?", pf.existe_execucao_fiscal),
            ("Existe Parcelamento?", pf.existe_parcelamento),
            ("Existe Transação Tributária?", pf.existe_transacao_tributaria),
            ("Existe discussão administrativa ou judicial?", pf.existe_discussao_administrativa_judicial),
        ],
        columns=["Pergunta", "Resposta"],
    )
    base.adicionar_tabela_dataframe(doc, df_situacao)

    doc.add_heading("Resumo", level=2)
    doc.add_paragraph(pf.resumo)

    doc.add_heading("Valores", level=2)
    df_valores = pd.DataFrame(
        [
            ("Valor do Passivo Fiscal", pf.valor_passivo_fiscal),
            ("Valor das Execuções Fiscais", pf.valor_execucoes_fiscais),
            ("Quantidade de Processos", pf.quantidade_processos),
            ("Tributos Envolvidos", ", ".join(pf.tributos_envolvidos) if pf.tributos_envolvidos else "Não localizado"),
            ("Órgãos Envolvidos", ", ".join(pf.orgaos_envolvidos) if pf.orgaos_envolvidos else "Não localizado"),
        ],
        columns=["Item", "Valor"],
    )
    base.adicionar_tabela_dataframe(doc, df_valores)

    doc.add_heading("Trechos Localizados", level=2)
    if pf.trechos_localizados:
        df_trechos = pd.DataFrame(
            [{"Página": t.pagina, "Trecho": t.trecho, "Contexto": t.contexto} for t in pf.trechos_localizados]
        )
        base.adicionar_tabela_dataframe(doc, df_trechos)
    else:
        doc.add_paragraph("Nenhum trecho localizado no documento.")

    doc.add_heading("Avaliação Estratégica", level=2)
    doc.add_paragraph(pf.avaliacao_estrategica or "Não foi possível gerar esta seção automaticamente.")

    doc.add_heading("Grau de Atenção", level=2)
    icone_grau = _ICONE_GRAU_ATENCAO.get(pf.grau_atencao, "🟢")
    paragrafo_grau = doc.add_paragraph()
    paragrafo_grau.add_run(f"{icone_grau} {pf.grau_atencao}").bold = True
    if pf.justificativa_grau_atencao:
        doc.add_paragraph(pf.justificativa_grau_atencao)


def _renderizar_secao(doc: Document, titulo: str, valor: object) -> None:
    doc.add_heading(titulo, level=1)

    if isinstance(valor, PassivoFiscal):
        _renderizar_passivo_fiscal(doc, valor)
        return

    if isinstance(valor, DadosEmpresa):
        _renderizar_dados_empresa(doc, valor)
        return

    if isinstance(valor, list):
        if not valor:
            doc.add_paragraph("Não foi possível identificar itens para esta seção no documento.")
        elif isinstance(valor[0], EventoCronologia):
            df = pd.DataFrame([{"Data": e.data, "Evento": e.evento} for e in valor])
            base.adicionar_tabela_dataframe(doc, df)
        elif isinstance(valor[0], ItemComJustificativa):
            df = pd.DataFrame([{"Ponto": i.ponto, "Justificativa": i.justificativa} for i in valor])
            base.adicionar_tabela_dataframe(doc, df)
        return

    doc.add_paragraph(str(valor) or "Não foi possível gerar esta seção automaticamente.")


def exportar_word_peticao_inicial(relatorio: RelatorioPeticaoInicial, caminho_saida: str | Path) -> Path:
    """Gera o relatório .docx completo (capa + sumário + 13 seções) a partir
    de um `RelatorioPeticaoInicial` já preenchido (ver `src/ia.py`).
    """
    caminho_saida = Path(caminho_saida)
    caminho_saida.parent.mkdir(parents=True, exist_ok=True)

    doc = Document()
    base.adicionar_capa(
        doc,
        relatorio.arquivo_nome,
        subtitulo_modulo="Relatório de Análise de Petição Inicial",
        texto_aviso=_TEXTO_AVISO_CAPA,
        pagina_isolada=True,
    )
    base.adicionar_sumario(doc)

    for titulo, chave in SECOES:
        _renderizar_secao(doc, titulo, getattr(relatorio, chave))

    if relatorio.paginas_ocr_baixa_confianca or relatorio.avisos:
        doc.add_heading("Avisos sobre a Extração/Geração", level=1)
        if relatorio.paginas_ocr_baixa_confianca:
            paginas = ", ".join(str(p) for p in relatorio.paginas_ocr_baixa_confianca)
            doc.add_paragraph(
                f"As páginas {paginas} foram lidas via OCR com possível baixa qualidade de "
                "reconhecimento — recomenda-se revisão manual dessas páginas no PDF original.",
                style="List Bullet",
            )
        for aviso in relatorio.avisos:
            doc.add_paragraph(aviso, style="List Bullet")

    base.adicionar_rodape_paginacao(doc, f"{NOME_EMPRESA} — Análise de Petição Inicial")

    doc.save(caminho_saida)
    logger.info("Word da Petição Inicial exportado em '%s' (%s).", caminho_saida, relatorio.arquivo_nome)
    return caminho_saida
