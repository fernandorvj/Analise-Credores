"""Helpers de marca compartilhados por todos os exportadores Word (.docx) do
sistema — capa, tabela a partir de DataFrame, sumário e rodapé paginado.

Extraído de `src/exportar_word.py` (módulo Credores) para ser reutilizado
pelo módulo Petição Inicial sem duplicar código. `exportar_word.py` continua
sendo o único responsável pelo conteúdo específico do relatório de Credores;
este módulo só sabe montar as peças genéricas de um documento com a
identidade visual da AMF3 Capital.
"""

from __future__ import annotations

from datetime import date

import pandas as pd
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Inches, Pt, RGBColor

from config import LOGO_PATH, NOME_EMPRESA, NOME_SISTEMA
from src.utils import formatar_moeda, formatar_percentual

COR_PRIMARIA_RGB = RGBColor(0x24, 0x22, 0x88)


def cor_hex_para_rgb(cor_hex: str) -> tuple[float, float, float]:
    cor_hex = cor_hex.lstrip("#")
    return tuple(int(cor_hex[i : i + 2], 16) / 255 for i in (0, 2, 4))


def adicionar_capa(
    doc: Document,
    nome_arquivo_pdf: str,
    subtitulo_modulo: str,
    texto_aviso: str,
    pagina_isolada: bool = False,
) -> None:
    """Capa padrão: logo, título do sistema, identificação do documento
    analisado e aviso legal. `pagina_isolada=True` força o conteúdo seguinte
    (ex.: sumário) para a próxima página — usado por relatórios mais longos,
    com várias seções, onde a capa não deve dividir página com conteúdo.
    """
    if LOGO_PATH.exists():
        doc.add_picture(str(LOGO_PATH), width=Inches(1.5))

    titulo = doc.add_heading(NOME_SISTEMA, level=0)
    titulo.runs[0].font.color.rgb = COR_PRIMARIA_RGB

    subtitulo = doc.add_paragraph()
    subtitulo.add_run(f"{NOME_EMPRESA} — {subtitulo_modulo}").bold = True
    subtitulo.add_run(f"\nDocumento analisado: {nome_arquivo_pdf}")
    subtitulo.add_run(f"\nData de geração: {date.today().strftime('%d/%m/%Y')}")

    aviso = doc.add_paragraph()
    run = aviso.add_run(texto_aviso)
    run.italic = True
    run.font.size = Pt(9)

    doc.add_paragraph()

    if pagina_isolada:
        doc.add_page_break()


def adicionar_tabela_dataframe(
    doc: Document,
    df: pd.DataFrame,
    colunas_moeda: set[str] = frozenset(),
    colunas_percentual: set[str] = frozenset(),
) -> None:
    if df.empty:
        doc.add_paragraph("Sem dados disponíveis para esta seção.")
        return

    tabela = doc.add_table(rows=1, cols=len(df.columns))
    tabela.style = "Light Grid Accent 1"
    for i, coluna in enumerate(df.columns):
        tabela.rows[0].cells[i].text = str(coluna)

    for _, linha_df in df.iterrows():
        celulas = tabela.add_row().cells
        for i, coluna in enumerate(df.columns):
            valor = linha_df[coluna]
            if coluna in colunas_moeda:
                texto = formatar_moeda(valor)
            elif coluna in colunas_percentual:
                texto = formatar_percentual(valor)
            else:
                texto = str(valor)
            celulas[i].text = texto


def _adicionar_campo(run, instrucao: str) -> None:
    """Insere um campo de campo do Word (ex.: PAGE, NUMPAGES, TOC) num run já
    existente, via o trio de elementos OOXML fldChar begin/instrText/end —
    a receita padrão do python-docx para campos que a própria biblioteca não
    calcula (ela não tem motor de layout, então não sabe o número da página).
    """
    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = instrucao
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_end)


def adicionar_sumario(doc: Document) -> None:
    """Campo de Sumário (TOC) apontando para os títulos "Heading 1" do
    documento. O Word recalcula o conteúdo real ao abrir o arquivo (setting
    `updateFields`); até lá, mostra um texto de espera pedindo para atualizar
    manually (Botão direito → Atualizar Campo), caso o Word não faça sozinho.
    """
    doc.add_heading("Sumário", level=1)
    paragrafo = doc.add_paragraph()
    run = paragrafo.add_run()

    fld_begin = OxmlElement("w:fldChar")
    fld_begin.set(qn("w:fldCharType"), "begin")
    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = 'TOC \\o "1-3" \\h \\z \\u'
    fld_separate = OxmlElement("w:fldChar")
    fld_separate.set(qn("w:fldCharType"), "separate")
    texto_placeholder = OxmlElement("w:t")
    texto_placeholder.text = (
        'Sumário — clique com o botão direito e escolha "Atualizar Campo" '
        "se não atualizar automaticamente ao abrir o documento."
    )
    fld_end = OxmlElement("w:fldChar")
    fld_end.set(qn("w:fldCharType"), "end")

    r_element = run._r
    r_element.append(fld_begin)
    r_element.append(instr)
    r_element.append(fld_separate)
    r_element.append(texto_placeholder)
    r_element.append(fld_end)

    # Faz o Word atualizar todos os campos (TOC, PAGE, NUMPAGES) ao abrir,
    # em vez de depender do usuário apertar F9 manualmente.
    settings = doc.settings.element
    update_fields = OxmlElement("w:updateFields")
    update_fields.set(qn("w:val"), "true")
    settings.append(update_fields)

    doc.add_page_break()


def adicionar_rodape_paginacao(doc: Document, texto_rodape: str) -> None:
    """Rodapé com o texto informado + numeração "Página X de Y" via campos
    PAGE/NUMPAGES (o Word calcula os números reais; python-docx não tem como
    calculá-los sozinho por não ter motor de layout).
    """
    footer = doc.sections[0].footer
    paragrafo = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
    paragrafo.alignment = WD_ALIGN_PARAGRAPH.CENTER

    run_texto = paragrafo.add_run(f"{texto_rodape} — Página ")
    run_texto.font.size = Pt(8)

    run_pagina = paragrafo.add_run()
    _adicionar_campo(run_pagina, "PAGE")
    run_pagina.font.size = Pt(8)

    run_de = paragrafo.add_run(" de ")
    run_de.font.size = Pt(8)

    run_total = paragrafo.add_run()
    _adicionar_campo(run_total, "NUMPAGES")
    run_total.font.size = Pt(8)
