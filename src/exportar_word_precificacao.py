"""Exportação para Word (.docx) do resultado da Precificação Inteligente de
Créditos — reaproveita a capa/sumário/rodapé compartilhados
(`src/exportar_word_base.py`), sem duplicá-los. Gráficos são renderizados
com `matplotlib` (mesmo padrão já usado em `src/calculadora/exportar_word.py`).
"""

from __future__ import annotations

import io
from pathlib import Path

import matplotlib

matplotlib.use("Agg")  # backend sem interface gráfica, necessário para servidores/Streamlit
import matplotlib.pyplot as plt
import pandas as pd
from docx import Document
from docx.shared import Inches

from config import NOME_EMPRESA
from src import exportar_word_base as base
from src.models_precificacao import ResultadoPrecificacaoClasse
from src.utils import formatar_moeda, formatar_percentual

_COR_GRAFICO = "#3633CC"
_COR_GRAFICO_2 = "#B87A12"

_TEXTO_AVISO_BASE = (
    "Esta análise combina condições de pagamento extraídas por Inteligência Artificial do Plano de "
    "Recuperação Judicial com cálculos financeiros determinísticos em Python — não constitui garantia "
    "de resultado, proposta de investimento ou aconselhamento jurídico/financeiro."
)


def _grafico_fluxo(resultado: ResultadoPrecificacaoClasse) -> io.BytesIO:
    fluxo_ordenado = sorted(resultado.fluxo, key=lambda item: item.numero)
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    datas = [p.data for p in fluxo_ordenado]
    nominal = [float(p.valor_nominal) for p in fluxo_ordenado]
    descontado = [float(p.valor_descontado) for p in fluxo_ordenado]
    ax.plot(datas, nominal, marker="o", color=_COR_GRAFICO, label="Valor Nominal")
    ax.plot(datas, descontado, marker="o", color=_COR_GRAFICO_2, label="Valor Presente (VP_t)")
    ax.set_title("Cronograma Unificado: Nominal x Valor Presente")
    ax.set_ylabel("Valor (R$)")
    plt.xticks(rotation=25, ha="right", fontsize=8)
    ax.legend()
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def _grafico_saldo(resultado: ResultadoPrecificacaoClasse) -> io.BytesIO:
    fluxo_ordenado = sorted(resultado.fluxo, key=lambda item: item.numero)
    fig, ax = plt.subplots(figsize=(6.5, 3.5))
    datas = [p.data for p in fluxo_ordenado]
    saldos = [float(p.saldo_final) for p in fluxo_ordenado]
    ax.plot(datas, saldos, color=_COR_GRAFICO, linewidth=2)
    ax.fill_between(datas, saldos, color=_COR_GRAFICO, alpha=0.12)
    ax.set_title("Evolução do Saldo Devedor")
    ax.set_ylabel("Saldo (R$)")
    plt.xticks(rotation=25, ha="right", fontsize=8)
    fig.tight_layout()
    buffer = io.BytesIO()
    fig.savefig(buffer, format="png", dpi=150)
    plt.close(fig)
    buffer.seek(0)
    return buffer


def exportar_word_precificacao(resultado: ResultadoPrecificacaoClasse, caminho_saida: str | Path) -> Path:
    """Gera o relatório Word da Precificação Inteligente: capa, sumário,
    condições consideradas, premissas, cronograma unificado completo,
    gráficos, resultados e memória de cálculo (auditoria)."""
    caminho_saida = Path(caminho_saida)
    c = resultado.condicoes
    doc = Document()

    texto_aviso = _TEXTO_AVISO_BASE
    if not resultado.metodologia_validada:
        texto_aviso += " Metodologia de cálculo ainda pendente de validação final contra a planilha oficial da AMF3 Capital."

    base.adicionar_capa(
        doc,
        nome_arquivo_pdf=f"Precificação — {resultado.classe}",
        subtitulo_modulo="Precificação Inteligente de Créditos",
        texto_aviso=texto_aviso,
        pagina_isolada=True,
    )
    base.adicionar_sumario(doc)

    if not resultado.metodologia_validada:
        aviso = doc.add_paragraph()
        run = aviso.add_run(
            "⚠ Metodologia de cronograma unificado, casamento de período e descapitalização linha por "
            "linha ainda não validada contra a planilha oficial de VPL da AMF3 Capital. Trate os números "
            "deste relatório como uma estimativa até a confirmação."
        )
        run.bold = True

    doc.add_heading("Condições de Pagamento Consideradas", level=1)
    df_condicoes = pd.DataFrame(
        [
            ("Classe", resultado.classe),
            ("Deságio", c.desagio),
            ("Carência", c.carencia),
            ("Índice de Correção Monetária", c.correcao_monetaria_indice),
            ("Juros", c.juros),
            ("Número de Parcelas", c.numero_parcelas),
            ("Periodicidade", c.periodicidade),
            ("Data da 1ª Parcela", c.data_primeira_parcela),
            ("Parcela Balão", c.parcela_balao),
        ],
        columns=["Condição", "Valor Considerado"],
    )
    base.adicionar_tabela_dataframe(doc, df_condicoes)

    if c.trechos_localizados:
        doc.add_heading("Trechos do Plano (Auditoria)", level=1)
        df_trechos = pd.DataFrame([{"Página": t.pagina, "Trecho": t.trecho, "Contexto": t.contexto} for t in c.trechos_localizados])
        base.adicionar_tabela_dataframe(doc, df_trechos)

    doc.add_heading("Dados da Operação", level=1)
    df_operacao = pd.DataFrame(
        [
            ("Valor Nominal do Crédito (C0)", formatar_moeda(float(resultado.valor_nominal_credito))),
            (
                "Valor Atualizado do Crédito",
                formatar_moeda(float(resultado.valor_atualizado_credito)) if resultado.valor_atualizado_credito is not None else "Não aplicável",
            ),
            ("Taxa de Desconto (a.a.)", formatar_percentual(float(resultado.taxa_desconto_anual))),
            ("Taxa de Desconto (por período)", formatar_percentual(float(resultado.taxa_desconto_periodo))),
            ("Origem da Taxa de Desconto", resultado.origem_taxa_desconto),
            (
                "Data da Taxa de Desconto",
                resultado.data_taxa_desconto.strftime("%d/%m/%Y") if resultado.data_taxa_desconto else "Manual",
            ),
        ],
        columns=["Parâmetro", "Valor"],
    )
    base.adicionar_tabela_dataframe(doc, df_operacao)

    doc.add_heading("Premissas", level=1)
    doc.add_paragraph(
        "Cronograma unificado: o número total de linhas é carência + número de parcelas — a carência entra "
        "na tabela como linhas próprias, sem pagamento ao credor (os juros do período capitalizam ao saldo "
        "devedor). Em toda linha, os juros incidem sobre o saldo devedor do período anterior, nunca sobre a "
        "parcela isolada. Casamento de período: a taxa de juros do plano e a taxa de desconto do VPL são "
        "sempre convertidas para a mesma periodicidade das parcelas antes de qualquer cálculo. "
        "Descapitalização linha por linha: VP_t = Fluxo Total da linha / (1 + taxa de desconto)^t, onde t é "
        "o número sequencial da linha cronológica — a mesma fórmula de VPL discreto (Excel NPV)."
    )

    doc.add_heading("Cronograma Unificado", level=1)
    df_fluxo = pd.DataFrame(
        [
            {
                "Nº": p.numero,
                "Data": p.data.strftime("%d/%m/%Y"),
                "Descrição": p.descricao,
                "Carência?": "Sim" if p.carencia else "Não",
                "Saldo Inicial": float(p.saldo_inicial),
                "Juros": float(p.juros_periodo),
                "Amortização": float(p.amortizacao),
                "Valor": float(p.valor_nominal),
                "Saldo Final": float(p.saldo_final),
                "Valor Presente": float(p.valor_descontado),
            }
            for p in sorted(resultado.fluxo, key=lambda item: item.numero)
        ]
    )
    base.adicionar_tabela_dataframe(
        doc, df_fluxo, colunas_moeda={"Saldo Inicial", "Juros", "Amortização", "Valor", "Saldo Final", "Valor Presente"}
    )

    doc.add_heading("Gráficos", level=1)
    doc.add_picture(_grafico_fluxo(resultado), width=Inches(6))
    doc.add_picture(_grafico_saldo(resultado), width=Inches(6))

    doc.add_heading("Resultados", level=1)
    df_resultados = pd.DataFrame(
        [
            ("Fluxo Nominal Total", formatar_moeda(float(resultado.fluxo_nominal_total))),
            ("Valor Presente do Fluxo (VP Total)", formatar_moeda(float(resultado.vp_total))),
            ("VPL Real Comercial (VP Total − C0)", formatar_moeda(float(resultado.vpl_comercial))),
            ("Percentual de Recuperação Efetiva", formatar_percentual(float(resultado.percentual_recuperacao_efetiva) / 100)),
        ],
        columns=["Indicador", "Valor"],
    )
    base.adicionar_tabela_dataframe(doc, df_resultados)
    aviso_negativo = doc.add_paragraph()
    aviso_negativo.add_run(
        "Como se trata de uma Recuperação Judicial com deságio e prazo longo, um VPL Real Comercial "
        "significativamente negativo é esperado — reflete a perda real de poder de compra do credor, "
        "não um erro de cálculo."
    ).italic = True

    doc.add_heading("Memória de Cálculo", level=1)
    for linha in resultado.memoria_calculo:
        doc.add_paragraph(linha, style="List Bullet")

    base.adicionar_rodape_paginacao(doc, f"{NOME_EMPRESA} — Precificação Inteligente de Créditos")
    doc.save(caminho_saida)
    return caminho_saida
